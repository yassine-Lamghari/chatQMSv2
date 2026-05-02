from fastapi import FastAPI, Depends, HTTPException, status, UploadFile, File, Form, BackgroundTasks, Query, Request, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from sqlalchemy.orm import Session
from database import (
    init_db,
    get_db,
    SessionLocal,
    LLMConfig,
    User,
    DocumentMetadata,
    DocumentTemplate,
    AppSetting,
    ActivityLog,
    ChatSession,
    AuditResult,
    seed_default_templates,
)
from crypto_utils import encrypt_api_key, decrypt_api_key  # Fix #4
from pydantic import BaseModel
import bcrypt
import os
import math
import json
import io
import logging
from datetime import datetime, timezone
from typing import Any, Optional
from auth import create_access_token, decode_access_token
from rag import (
    ingest_document,
    remove_document_from_index,
    search_similar_chunks,
    SUPPORTED_EXTENSIONS,
    EMBEDDING_MODEL_NAME,
    CHROMA_PERSIST_DIR,
)

logger = logging.getLogger(__name__)
from services_qms import (
    audit_questions_for_standard,
    audit_sampling_plan,
    pfmea_skeleton_rows,
    verify_pfmea_row,
)
from llm_rag import build_numbered_context, synthesize_from_context, generate_pfmea_rows_llm

limiter = Limiter(key_func=get_remote_address)

os.makedirs("uploads", exist_ok=True)


def _log_activity(
    db: Session,
    action: str,
    username: str = "anonymous",
    query: str | None = None,
    document_ids: list[str] | None = None,
    confidence: str | None = None,
    language_mode: str | None = None,
    response_summary: str | None = None,
):
    """Insert an activity log entry (non-blocking best-effort)."""
    try:
        entry = ActivityLog(
            username=username,
            action=action,
            query=(query or "")[:500],
            document_ids=",".join(document_ids) if document_ids else None,
            confidence=confidence,
            language_mode=language_mode,
            response_summary=(response_summary or "")[:300],
        )
        db.add(entry)
        db.commit()
    except Exception as _e:
        logger.warning("ActivityLog write failed: %s", _e)
        db.rollback()

def verify_password(plain_password: str, hashed_password: str) -> bool:
    try:
        return bcrypt.checkpw(plain_password.encode('utf-8'), hashed_password.encode('utf-8'))
    except ValueError:
        return False

def get_password_hash(password: str) -> str:
    salt = bcrypt.gensalt()
    return bcrypt.hashpw(password.encode('utf-8'), salt).decode('utf-8')

class UserCreate(BaseModel):
    username: str
    password: str
    role: str = "user"

class UserLogin(BaseModel):
    username: str
    password: str

class ChatRequest(BaseModel):
    query: str
    top_k: int = 5
    language_mode: str = "document_language"  # en_only | document_language | fr_with_en_sources
    respond_in_english: bool = False
    response_locale: str = "fr"  # fr | en (UI / canned messages)
    user_role: str = "user"  # user | admin — used for criticality access rules
    filters: dict = {}
    # When True (default), call the active LLM to synthesize summary/details from retrieved chunks.
    use_llm: bool = True
    username: str = "anonymous"   # passed from frontend for activity logging

class ChatSource(BaseModel):
    filename: str
    doc_type: str
    criticality: str
    doc_id: str
    language: str
    version: str
    owner: str
    relevance: float

class ChatResponse(BaseModel):
    summary: str
    summary_bullets: list[str]
    details: str
    confidence: str
    sources: list[dict]

class AuditChecklistRequest(BaseModel):
    standard: str = "ISO 9001"
    process: str
    top_k: int = 5

class SharePointSyncRequest(BaseModel):
    site_url: str
    library_name: str


class SharePointConfigRequest(BaseModel):
    tenant_id: str = ""
    client_id: str = ""
    client_secret: str = ""
    site_url: str = ""
    library_id: str = ""


class SearchRequest(BaseModel):
    query: str
    top_k: int = 8
    filters: dict = {}


class PfmeaGenerateRequest(BaseModel):
    process: str
    product: str
    known_defects: str = ""
    top_k: int = 4


class VerifyDocumentRequest(BaseModel):
    mode: str = "pfmea_row"  # pfmea_row
    data: dict[str, Any]


class AuditAssistantRequest(BaseModel):
    standard: str = "ISO 9001"
    process: str
    depth: str = "normal"
    top_k: int = 5

class ActiveLLMRequest(BaseModel):
    provider: str

class LLMProviderConfigRequest(BaseModel):
    api_key: str | None = None
    base_url: str | None = None

app = FastAPI(title="QMS Chatbot API")
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

MAX_DISTANCE_THRESHOLD = 2.35
MIN_RELEVANCE_THRESHOLD = 0.12

_ALLOWED_ORIGINS = os.getenv("FRONTEND_URL", "http://localhost:3000").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=_ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Dependency: optional JWT auth (non-blocking for existing endpoints) ──
def get_optional_user(authorization: Optional[str] = Header(None)) -> dict | None:
    if not authorization or not authorization.startswith("Bearer "):
        return None
    token = authorization[7:]
    return decode_access_token(token)

def require_auth(authorization: Optional[str] = Header(None)) -> dict:
    user = get_optional_user(authorization)
    if not user:
        raise HTTPException(status_code=401, detail="Authentication required")
    return user

def require_admin(authorization: Optional[str] = Header(None)) -> dict:
    user = require_auth(authorization)
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return user

@app.on_event("startup")
def on_startup():
    init_db()
    db = SessionLocal()
    try:
        seed_default_templates(db)
    finally:
        db.close()
    logger.warning(
        "RAG: embedding model is %s (Chroma dir=%s). If results are wrong or empty after an upgrade, "
        "delete the Chroma folder and re-upload documents so vectors are rebuilt.",
        EMBEDDING_MODEL_NAME,
        CHROMA_PERSIST_DIR,
    )
    # Pré-chargement du modèle d'embedding au démarrage pour éviter le délai
    # sur la première requête utilisateur (le modèle peut prendre 1-2 min à charger).
    try:
        from rag import get_embeddings, get_vector_store
        logger.warning("RAG: pré-chargement du modèle d'embedding en mémoire...")
        get_embeddings()
        get_vector_store()
        logger.warning("RAG: modèle pret !")
    except Exception as _e:
        logger.error("RAG: echec du pre-chargement -- %s", _e)


def _section_ref(doc) -> str:
    meta = doc.metadata or {}
    fn = meta.get("filename", "document")
    page = meta.get("page")
    pl = meta.get("page_label")
    if page is not None:
        return f"{fn} — § page {int(page) + 1 if isinstance(page, int) and page < 500 else page}" + (f" (label {pl})" if pl else "")
    return f"{fn} — section (page metadata n/a)"


def _can_access_criticality(user_role: str, criticality: str) -> bool:
    if user_role == "admin":
        return True
    c = (criticality or "").strip().lower()
    if c == "critical":
        return False
    return True


def _confidence_label(score: float, locale: str) -> str:
    if (locale or "fr").lower().startswith("en"):
        return "High" if score >= 0.7 else "Medium" if score >= 0.45 else "Low"
    return "Élevé" if score >= 0.7 else "Moyen" if score >= 0.45 else "Faible"


def _active_llm_settings(db: Session) -> tuple[str | None, str | None, str | None, str | None]:
    """
    Active provider plus credentials for chat synthesis.
    Returns (provider, api_key, base_url, ollama_model_name).
    """
    row = db.query(AppSetting).filter(AppSetting.key == "active_llm_provider").first()
    provider = (row.value or "").strip().lower() if row and row.value else None
    if not provider:
        return None, None, None, None
    if provider == "ollama":
        cfg = db.query(LLMConfig).filter(LLMConfig.provider == "ollama").first()
        base = (cfg.base_url.strip() if cfg and cfg.base_url else "") or os.getenv("OLLAMA_BASE_URL", "http://127.0.0.1:11434")
        model = os.getenv("OLLAMA_MODEL", "llama3.2")
        return provider, None, base, model
    cfg = db.query(LLMConfig).filter(LLMConfig.provider == provider).first()
    if not cfg:
        return provider, None, None, None
    key = decrypt_api_key((cfg.api_key or "").strip()) or None  # Fix #4
    base = (cfg.base_url or "").strip() or None
    return provider, key, base, None


def _not_found_payload(locale: str):
    loc = (locale or "fr").lower()
    if loc.startswith("en"):
        return {
            "summary": "No relevant context found in the document base.",
            "summary_bullets": ["- No indexed source matched this query."],
            "details": "Upload QMS documents (PDF/Word) from the admin panel, then try again.",
            "detail_sections": [],
            "confidence": "Low",
            "sources": [],
        }
    return {
        "summary": "Aucun contexte pertinent trouvé dans la base documentaire.",
        "summary_bullets": ["- Aucune source indexée ne correspond à cette requête."],
        "details": "Importez des documents QMS (PDF/Word) depuis l'administration, puis réessayez.",
        "detail_sections": [],
        "confidence": "Faible",
        "sources": [],
    }

@app.get("/")
def read_root():
    return {"message": "Welcome to the QMS Chatbot API"}


# Fix #17 — Indicateur de statut LLM pour l'UI
@app.get("/api/llm/status")
def get_llm_status(db: Session = Depends(get_db)):
    """Retourne si un LLM est configuré et opérationnel."""
    provider, api_key, base_url, ollama_model = _active_llm_settings(db)
    if not provider:
        return {"configured": False, "provider": None, "message": "Aucun LLM configuré — résultats en mode extrait uniquement"}
    if provider == "ollama":
        # Vérifier que Ollama est accessible
        import urllib.request, urllib.error
        try:
            with urllib.request.urlopen(f"{base_url}/api/tags", timeout=2):
                return {"configured": True, "provider": provider, "message": f"Ollama actif ({base_url})"}
        except Exception:
            return {"configured": False, "provider": provider, "message": f"Ollama non joignable sur {base_url}"}
    # Provider cloud : vérifier qu'une clé est présente
    if api_key and len(api_key) > 5:
        return {"configured": True, "provider": provider, "message": f"LLM actif : {provider}"}
    return {"configured": False, "provider": provider, "message": f"Clé API manquante pour {provider}"}

@app.get("/api/config")
def get_llm_configs(db: Session = Depends(get_db), _user: dict = Depends(require_admin)):  # Fix #2
    configs = db.query(LLMConfig).all()
    # Fix #4: ne jamais retourner les clés en clair
    return [
        {"id": c.id, "provider": c.provider, "api_key": "***" if c.api_key else None, "base_url": c.base_url, "model_name": c.model_name}
        for c in configs
    ]

@app.post("/api/config")
def update_llm_config(provider: str, api_key: str = None, base_url: str = None, db: Session = Depends(get_db), _user: dict = Depends(require_admin)):  # Fix #2
    config = db.query(LLMConfig).filter(LLMConfig.provider == provider).first()
    if not config:
        config = LLMConfig(provider=provider, api_key=encrypt_api_key(api_key), base_url=base_url)  # Fix #4
        db.add(config)
    else:
        if api_key is not None:
            config.api_key = encrypt_api_key(api_key)  # Fix #4
        if base_url is not None:
            config.base_url = base_url
    db.commit()
    db.refresh(config)
    return {"provider": config.provider, "base_url": config.base_url}

@app.get("/api/config/active")
def get_active_llm(db: Session = Depends(get_db), _user: dict = Depends(require_admin)):  # Fix #2
    setting = db.query(AppSetting).filter(AppSetting.key == "active_llm_provider").first()
    return {"provider": setting.value if setting else None}


@app.post("/api/config/active")
def set_active_llm(payload: ActiveLLMRequest, db: Session = Depends(get_db), _user: dict = Depends(require_admin)):  # Fix #2
    provider = payload.provider.strip().lower()
    config = db.query(LLMConfig).filter(LLMConfig.provider == provider).first()
    if not config:
        config = LLMConfig(provider=provider, api_key=None, base_url=None)
        db.add(config)

    setting = db.query(AppSetting).filter(AppSetting.key == "active_llm_provider").first()
    if not setting:
        setting = AppSetting(key="active_llm_provider", value=provider)
        db.add(setting)
    else:
        setting.value = provider
    db.commit()
    return {"provider": provider, "message": "Active LLM updated"}


@app.get("/api/config/{provider}")
def get_llm_provider_config(provider: str, db: Session = Depends(get_db)):
    if provider.lower() == "active":
        raise HTTPException(status_code=404, detail="Not found")
    config = db.query(LLMConfig).filter(LLMConfig.provider == provider.lower()).first()
    if not config:
        raise HTTPException(status_code=404, detail="Provider not found")
    return config


@app.put("/api/config/{provider}")
def upsert_llm_provider_config(provider: str, payload: LLMProviderConfigRequest, db: Session = Depends(get_db), _user: dict = Depends(require_admin)):  # Fix #2
    provider_name = provider.strip().lower()
    if provider_name == "active":
        raise HTTPException(status_code=400, detail="Invalid provider name")
    if provider_name == "ollama":
        raise HTTPException(status_code=400, detail="Use dedicated local settings for ollama")

    config = db.query(LLMConfig).filter(LLMConfig.provider == provider_name).first()
    if not config:
        config = LLMConfig(provider=provider_name)
        db.add(config)

    if payload.api_key is not None:
        config.api_key = encrypt_api_key(payload.api_key.strip()) if payload.api_key else None  # Fix #4
    if payload.base_url is not None:
        config.base_url = payload.base_url.strip() if payload.base_url else None

    db.commit()
    db.refresh(config)
    return {"provider": config.provider, "base_url": config.base_url}

@app.post("/api/chat")
def chat(payload: ChatRequest, db: Session = Depends(get_db)):
    query = payload.query.strip()
    if not query:
        raise HTTPException(status_code=400, detail="Query cannot be empty")
    top_k = max(1, min(payload.top_k, 8))
    locale = (payload.response_locale or "fr").lower()
    user_role = (payload.user_role or "user").strip().lower()
    language_mode = payload.language_mode.strip()
    if language_mode not in {"en_only", "document_language", "fr_with_en_sources"}:
        raise HTTPException(status_code=400, detail="Invalid language_mode")

    raw_filters = payload.filters or {}
    metadata_filter: dict[str, str] = {}
    for key in ("doc_type", "criticality", "language", "owner", "version", "site"):
        value = raw_filters.get(key)
        if isinstance(value, str) and value.strip():
            metadata_filter[key] = value.strip()

    date_from = raw_filters.get("date_from")
    date_to = raw_filters.get("date_to")

    retrievals = search_similar_chunks(
        query=query,
        k=min(top_k * 2, 16),
        metadata_filter=metadata_filter if metadata_filter else None,
    )

    if not retrievals:
        return _not_found_payload(locale)

    rag_hits = [(doc, float(distance)) for doc, distance in retrievals if float(distance) <= MAX_DISTANCE_THRESHOLD]
    if not rag_hits:
        nf = _not_found_payload(locale)
        nf["summary"] = (
            "RAG only: no passage met the relevance threshold."
            if locale.startswith("en")
            else "RAG only: aucun passage ne dépasse le seuil de pertinence."
        )
        nf["summary_bullets"] = [
            "- No source passed relevance threshold." if locale.startswith("en") else "- Aucune source ne dépasse le seuil de pertinence."
        ]
        return nf

    def _parse_dt(v):
        if not v or not isinstance(v, str):
            return None
        try:
            return datetime.fromisoformat(v.replace("Z", "+00:00")).replace(tzinfo=None)
        except ValueError:
            return None

    df = _parse_dt(date_from) if date_from else None
    dt = _parse_dt(date_to) if date_to else None

    filtered = []
    for doc, dist in rag_hits:
        meta = doc.metadata or {}
        source_doc_id = str(meta.get("doc_id", "unknown"))
        db_doc = db.query(DocumentMetadata).filter(DocumentMetadata.id == int(source_doc_id)).first() if source_doc_id.isdigit() else None
        crit = (meta.get("criticality") or (db_doc.criticality if db_doc else "") or "").strip()
        if not _can_access_criticality(user_role, crit):
            continue
        uploaded_at = db_doc.uploaded_at if db_doc else None
        if df and uploaded_at and uploaded_at < df:
            continue
        if dt and uploaded_at and uploaded_at > dt:
            continue
        filtered.append((doc, dist))

    rag_hits = filtered
    if not rag_hits:
        return _not_found_payload(locale)

    # Language filtering is intentionally NOT done at the RAG/vector-search level.
    # The LLM handles the response language via its system prompt (respond_english flag).
    # Filtering by language tag here would silently block relevant chunks from
    # documents that were uploaded without an explicit language tag.

    sources = []
    context_snippets = []
    distance_values = []
    freshness_scores = []
    detail_sections = []
    for doc, score in rag_hits[:10]:
        distance_values.append(score)
        meta = doc.metadata or {}
        source_doc_id = str(meta.get("doc_id", "unknown"))
        db_doc = db.query(DocumentMetadata).filter(DocumentMetadata.id == int(source_doc_id)).first() if source_doc_id.isdigit() else None
        uploaded_at = db_doc.uploaded_at if db_doc else None
        freshness = 0.5
        if uploaded_at:
            age_days = max((datetime.now(timezone.utc).replace(tzinfo=None) - uploaded_at).days, 0)
            freshness = max(0.1, 1 - (age_days / 365))
        freshness_scores.append(freshness)

        site_val = meta.get("site") or (db_doc.site if db_doc else "default")
        source = {
            "filename": meta.get("filename", "unknown"),
            "doc_type": meta.get("doc_type", "unknown"),
            "criticality": meta.get("criticality") or (db_doc.criticality if db_doc else "unknown"),
            "doc_id": source_doc_id,
            "language": meta.get("language") or (db_doc.language if db_doc else "unknown"),
            "version": meta.get("version") or (db_doc.version if db_doc else "unknown"),
            "owner": meta.get("owner") or (db_doc.owner if db_doc else "unknown"),
            "site": site_val,
            "relevance": round(math.exp(-score), 4),
            "section_ref": _section_ref(doc),
        }
        sources.append(source)
        excerpt = doc.page_content.strip()
        context_snippets.append(excerpt)
        detail_sections.append(
            {
                "section_ref": source["section_ref"],
                "excerpt": excerpt[:4000],
                "filename": source["filename"],
                "page": meta.get("page"),
                "doc_id": source_doc_id,
                "language": source["language"],
            }
        )

    unique_sources = []
    seen = set()
    for source in sources:
        key = (source["doc_id"], source["filename"], source["version"], source["section_ref"])
        if key not in seen:
            seen.add(key)
            unique_sources.append(source)

    excerpt_summary_bullets: list[str] = []
    for i, snippet in enumerate(context_snippets[:10]):
        line = snippet.replace("\n", " ").strip()
        if not line:
            continue
        prefix = ""
        if language_mode == "fr_with_en_sources" and not locale.startswith("en"):
            lang = (sources[i]["language"] if i < len(sources) else "").lower()
            if lang.startswith("en"):
                prefix = "[Source EN] "
        excerpt_summary_bullets.append(f"- {prefix}{line[:160]}")
    if len(excerpt_summary_bullets) < 5:
        for s in unique_sources[: 10 - len(excerpt_summary_bullets)]:
            excerpt_summary_bullets.append(f"- ({s['section_ref']}) relevance={s['relevance']}")
    excerpt_summary_bullets = excerpt_summary_bullets[:10]

    relevance_values = [math.exp(-d) for d in distance_values]
    avg_relevance = sum(relevance_values) / max(len(relevance_values), 1)
    best_relevance = max(relevance_values) if relevance_values else 0.0
    if best_relevance < MIN_RELEVANCE_THRESHOLD:
        return {
            "summary": (
                "RAG only: not found in the document base (weak match)."
                if locale.startswith("en")
                else "RAG only: information non trouvée (correspondance trop faible)."
            ),
            "summary_bullets": [
                "- The question does not align enough with indexed documents."
                if locale.startswith("en")
                else "- La question ne correspond pas assez aux documents indexés."
            ],
            "details": (
                "Rephrase the question or add a document covering this topic."
                if locale.startswith("en")
                else "Reformulez la question ou ajoutez un document couvrant ce sujet."
            ),
            "detail_sections": [],
            "confidence": "Low" if locale.startswith("en") else "Faible",
            "sources": [],
        }

    source_coverage = min(len(unique_sources) / 3, 1.0)
    freshness_score = sum(freshness_scores) / max(len(freshness_scores), 1)
    confidence_score = (0.5 * avg_relevance) + (0.3 * source_coverage) + (0.2 * freshness_score)
    confidence = _confidence_label(confidence_score, locale)

    combined_context = "\n\n---\n\n".join(context_snippets[:3])
    detail_intro = (
        "Detailed excerpts with source references (Vue 2):"
        if locale.startswith("en")
        else "Détail avec renvoi aux sections sources (Vue 2) :"
    )
    if language_mode == "fr_with_en_sources" and not locale.startswith("en"):
        combined_context = (
            "Résumé en français ; les extraits ci-dessous reprennent le texte source (souvent en anglais) tel qu’indexé.\n\n"
            + combined_context
        )

    # respond_english drives the LLM system prompt language.
    # en_only  → always EN
    # fr_with_en_sources → always FR (LLM prompt in FR but cites EN excerpts)
    # document_language  → let the LLM infer from context (defaults to FR unless forced)
    respond_english = bool(
        payload.respond_in_english or language_mode == "en_only"
    )
    generation_mode = "excerpts"
    rag_synthesis: str | None = None
    summary_text = (
        f"Answer based strictly on indexed documents for: {query}"
        if respond_english
        else f"Réponse basée uniquement sur la base documentaire pour: {query}"
    )
    summary_bullets = list(excerpt_summary_bullets)

    active_p, api_k, base_u, ollama_m = _active_llm_settings(db)
    can_llm = bool(
        payload.use_llm
        and active_p
        and (active_p == "ollama" or (api_k and len(api_k) > 0))
    )
    if can_llm:
        refs = [s["section_ref"] for s in sources[:top_k]]
        snips = context_snippets[:top_k]
        numbered = build_numbered_context(refs, snips, top_k) if refs and snips else ""
        if numbered.strip():
            try:
                syn = synthesize_from_context(
                    provider=active_p,
                    api_key=api_k,
                    base_url=base_u,
                    ollama_model=ollama_m,
                    user_query=query,
                    numbered_context=numbered,
                    respond_english=respond_english,
                    language_mode=language_mode,
                )
                if syn:
                    generation_mode = "llm"
                    summary_text = syn["summary"]
                    bullets = syn.get("summary_bullets") or []
                    summary_bullets = []
                    for b in bullets:
                        b = (b or "").strip()
                        if not b:
                            continue
                        summary_bullets.append(b if b.startswith("-") else f"- {b}")
                    if not summary_bullets:
                        summary_bullets = [
                            f"- {summary_text[:200]}..." if len(summary_text) > 200 else f"- {summary_text}"
                        ]
                    rag_synthesis = (syn.get("details") or "").strip() or None
                    if syn.get("answer_in_context") is False:
                        confidence = "Low" if respond_english else "Faible"
            except Exception as e:
                logger.warning("LLM RAG synthesis failed (%s): %s", active_p, e)

    details_body = f"{detail_intro}\n\n{combined_context}" if combined_context else detail_intro
    if rag_synthesis:
        syn_intro = (
            "Synthesized answer (grounded on the passages below):"
            if respond_english
            else "Synthèse (basée sur les extraits ci-dessous) :"
        )
        details_body = f"{syn_intro}\n\n{rag_synthesis}\n\n---\n\n{detail_intro}\n\n{combined_context}"

    result = {
        "summary": summary_text,
        "summary_bullets": summary_bullets,
        "details": details_body,
        "detail_sections": detail_sections[:8],
        "confidence": confidence,
        "confidence_score": round(confidence_score, 3),
        "sources": unique_sources,
        "rag_synthesis": rag_synthesis,
        "generation_mode": generation_mode,
    }

    _log_activity(
        db,
        action="chat",
        username=payload.username,
        query=query,
        document_ids=[s["doc_id"] for s in unique_sources],
        confidence=confidence,
        language_mode=language_mode,
        response_summary=summary_text,
    )

    return result

@app.post("/api/search")
def semantic_search(payload: SearchRequest, db: Session = Depends(get_db), _user: dict = Depends(require_auth)):  # Fix #3
    raw_filters = payload.filters or {}
    metadata_filter = {}
    for key in ("doc_type", "criticality", "language", "owner", "version", "site"):
        value = raw_filters.get(key)
        if isinstance(value, str) and value.strip():
            metadata_filter[key] = value.strip()
    k = max(1, min(payload.top_k, 20))
    hits = search_similar_chunks(
        query=payload.query.strip(),
        k=k,
        metadata_filter=metadata_filter if metadata_filter else None,
    )
    out = []
    for doc, distance in hits:
        if float(distance) > MAX_DISTANCE_THRESHOLD:
            continue
        rel = round(math.exp(-float(distance)), 4)
        if rel < MIN_RELEVANCE_THRESHOLD:
            continue
        meta = doc.metadata or {}
        source_doc_id = str(meta.get("doc_id", "unknown"))
        db_doc = db.query(DocumentMetadata).filter(DocumentMetadata.id == int(source_doc_id)).first() if source_doc_id.isdigit() else None
        crit = meta.get("criticality") or (db_doc.criticality if db_doc else "")
        out.append(
            {
                "section_ref": _section_ref(doc),
                "excerpt": doc.page_content.strip()[:800],
                "distance": float(distance),
                "relevance": rel,
                "filename": meta.get("filename"),
                "doc_id": source_doc_id,
                "doc_type": meta.get("doc_type"),
                "criticality": crit,
                "site": meta.get("site") or (db_doc.site if db_doc else None),
                "language": meta.get("language") or (db_doc.language if db_doc else None),
            }
        )
    return {"query": payload.query, "hits": out}


# Fix #15 — Recherche avec filtres étendus
class SearchRequestExtended(BaseModel):
    query: str
    top_k: int = 8
    filters: dict = {}
    # Nouveaux filtres
    criticality: str | None = None
    language: str | None = None
    owner: str | None = None
    date_from: str | None = None
    date_to: str | None = None

@app.post("/api/search/advanced")
def semantic_search_advanced(payload: SearchRequestExtended, db: Session = Depends(get_db), _user: dict = Depends(require_auth)):
    """Recherche sémantique avec filtres étendus (criticality, language, owner, dates)."""
    raw_filters = payload.filters or {}
    metadata_filter = {}
    for key in ("doc_type", "criticality", "language", "owner", "version", "site"):
        val = raw_filters.get(key) or getattr(payload, key, None)
        if isinstance(val, str) and val.strip():
            metadata_filter[key] = val.strip()

    k = max(1, min(payload.top_k, 20))
    hits = search_similar_chunks(
        query=payload.query.strip(),
        k=k,
        metadata_filter=metadata_filter if metadata_filter else None,
    )

    def _parse_dt(v):
        if not v or not isinstance(v, str):
            return None
        try:
            return datetime.fromisoformat(v.replace("Z", "+00:00")).replace(tzinfo=None)
        except ValueError:
            return None

    df = _parse_dt(payload.date_from)
    dt = _parse_dt(payload.date_to)

    out = []
    for doc, distance in hits:
        if float(distance) > MAX_DISTANCE_THRESHOLD:
            continue
        rel = round(math.exp(-float(distance)), 4)
        if rel < MIN_RELEVANCE_THRESHOLD:
            continue
        meta = doc.metadata or {}
        source_doc_id = str(meta.get("doc_id", "unknown"))
        db_doc = db.query(DocumentMetadata).filter(DocumentMetadata.id == int(source_doc_id)).first() if source_doc_id.isdigit() else None
        # Filtre date
        if db_doc and df and db_doc.uploaded_at and db_doc.uploaded_at < df:
            continue
        if db_doc and dt and db_doc.uploaded_at and db_doc.uploaded_at > dt:
            continue
        crit = meta.get("criticality") or (db_doc.criticality if db_doc else "")
        out.append({
            "section_ref": _section_ref(doc),
            "excerpt": doc.page_content.strip()[:800],
            "distance": float(distance),
            "relevance": rel,
            "filename": meta.get("filename"),
            "doc_id": source_doc_id,
            "doc_type": meta.get("doc_type"),
            "criticality": crit,
            "owner": meta.get("owner") or (db_doc.owner if db_doc else None),
            "language": meta.get("language") or (db_doc.language if db_doc else None),
            "site": meta.get("site") or (db_doc.site if db_doc else None),
            "version": meta.get("version") or (db_doc.version if db_doc else None),
            "uploaded_at": db_doc.uploaded_at.isoformat() if db_doc and db_doc.uploaded_at else None,
        })
    return {"query": payload.query, "hits": out, "total": len(out)}


@app.get("/api/templates")
def list_templates(db: Session = Depends(get_db)):
    return [
        {
            "id": t.id,
            "key": t.key,
            "name": t.name,
            "doc_type": t.doc_type,
            "language": t.language,
            "version": t.version,
        }
        for t in db.query(DocumentTemplate).all()
    ]


@app.post("/api/generate/pfmea")
def generate_pfmea(payload: PfmeaGenerateRequest, db: Session = Depends(get_db)):
    q = f"PFMEA failure mode process {payload.process} product {payload.product} defects {payload.known_defects}"
    hits = search_similar_chunks(query=q, k=max(1, min(payload.top_k, 8)))
    excerpts = [d.page_content.strip() for d, _ in hits if d.page_content]

    # Try LLM enrichment
    active_p, api_k, base_u, ollama_m = _active_llm_settings(db)
    can_llm = bool(active_p and (active_p == "ollama" or (api_k and len(api_k) > 0)))

    rows = None
    if can_llm:
        refs = [f"Source {i+1}" for i in range(len(hits))]
        numbered = build_numbered_context(refs, excerpts, payload.top_k)
        rows = generate_pfmea_rows_llm(
            provider=active_p,
            api_key=api_k,
            base_url=base_u,
            ollama_model=ollama_m,
            process=payload.process,
            product=payload.product,
            known_defects=payload.known_defects,
            numbered_context=numbered,
            respond_english=False,
        )

    if not rows:
        rows = pfmea_skeleton_rows(payload.process, payload.product, payload.known_defects, excerpts[:3])

    tpl = db.query(DocumentTemplate).filter(DocumentTemplate.key == "pfmea_blank").first()
    return {
        "template": tpl.key if tpl else None,
        "rows": rows,
        "rag_excerpts_used": excerpts[:payload.top_k],
    }


@app.post("/api/generate/verify")
def verify_document(payload: VerifyDocumentRequest):
    if payload.mode != "pfmea_row":
        raise HTTPException(status_code=400, detail="Unsupported verify mode")
    missing, warnings = verify_pfmea_row(payload.data)
    return {"missing_fields": missing, "warnings": warnings, "ok": len(missing) == 0}


@app.post("/api/audit/assistant")
def audit_assistant(payload: AuditAssistantRequest, db: Session = Depends(get_db)):
    std = payload.standard.strip()
    process = payload.process.strip()
    questions = audit_questions_for_standard(std, process)
    sampling = audit_sampling_plan(process, payload.depth)
    audit_plan = [
        {"day": 1, "focus": "Opening meeting, context, leadership, customer focus", "process": process},
        {"day": 2, "focus": "Operation & support processes, documented information", "process": process},
        {"day": 3, "focus": "Performance, improvement, audit closing", "process": process},
    ]
    retrievals = search_similar_chunks(
        query=f"{std} {process} audit checklist nonconformity",
        k=max(1, min(payload.top_k, 8)),
    )
    rag_checks = []
    sources = []
    for idx, (doc, distance) in enumerate(retrievals[:8], start=1):
        excerpt = doc.page_content.replace("\n", " ").strip()
        if excerpt and float(distance) <= MAX_DISTANCE_THRESHOLD:
            rel = round(math.exp(-float(distance)), 4)
            if rel >= MIN_RELEVANCE_THRESHOLD:
                rag_checks.append(f"{idx}. Evidence check: {excerpt[:160]}")
                sources.append(
                    {
                        "section_ref": _section_ref(doc),
                        "filename": doc.metadata.get("filename", "unknown"),
                        "doc_id": doc.metadata.get("doc_id", "unknown"),
                        "relevance": rel,
                    }
                )
    return {
        "standard": std,
        "process": process,
        "checklist_normative": questions,
        "audit_plan": audit_plan,
        "sampling": sampling,
        "rag_evidence_checks": rag_checks,
        "sources": sources,
    }


@app.get("/api/audit/export")
def export_audit_pack(
    standard: str = "ISO 9001",
    process: str = "Document control",
    export_format: str = Query("docx", alias="format"),
):
    """MVP export: Word (.docx) or simple PDF. Data is deterministic (no live RAG in export)."""
    questions = audit_questions_for_standard(standard, process)
    sampling = audit_sampling_plan(process)
    fmt = (export_format or "docx").lower()
    if fmt == "docx":
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx is not installed on the server")
        doc = DocxDocument()
        doc.add_heading("QMS Audit pack", level=0)
        doc.add_paragraph(f"Standard: {standard}")
        doc.add_paragraph(f"Process: {process}")
        doc.add_heading("Sampling", level=1)
        for k, v in sampling.items():
            doc.add_paragraph(f"{k}: {v}")
        doc.add_heading("Checklist", level=1)
        for q in questions:
            doc.add_paragraph(q, style="List Bullet")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="audit_pack_{process[:20]}.docx"'},
        )
    if fmt == "pdf":
        try:
            from fpdf import FPDF
        except ImportError:
            raise HTTPException(status_code=500, detail="fpdf2 is not installed on the server")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 8, f"QMS Audit pack\nStandard: {standard}\nProcess: {process}\n")
        pdf.ln(4)
        for k, v in sampling.items():
            pdf.multi_cell(0, 8, f"{k}: {v}")
        pdf.ln(4)
        for q in questions:
            pdf.multi_cell(0, 8, f"- {q}")
        raw = pdf.output()
        if isinstance(raw, str):
            raw = raw.encode("latin-1")
        out = io.BytesIO(raw)
        out.seek(0)
        return StreamingResponse(
            out,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="audit_pack_{process[:20]}.pdf"'},
        )
    raise HTTPException(status_code=400, detail="format must be docx or pdf")


@app.post("/api/audit/checklist")
def generate_audit_checklist(payload: AuditChecklistRequest, db: Session = Depends(get_db)):
    return audit_assistant(
        AuditAssistantRequest(standard=payload.standard, process=payload.process, depth="normal", top_k=payload.top_k),
        db,
    )


@app.post("/api/sharepoint/config")
def sharepoint_configure(payload: SharePointConfigRequest, db: Session = Depends(get_db)):
    data = {
        "tenant_id": payload.tenant_id,
        "client_id": payload.client_id,
        "client_secret_set": bool(payload.client_secret and payload.client_secret.strip()),
        "site_url": payload.site_url,
        "library_id": payload.library_id,
    }
    row = db.query(AppSetting).filter(AppSetting.key == "sharepoint_config").first()
    if not row:
        row = AppSetting(key="sharepoint_config", value=json.dumps(data))
        db.add(row)
    else:
        row.value = json.dumps(data)
    db.commit()
    return {"message": "SharePoint configuration stored (secrets are not echoed).", "site_url": payload.site_url}


@app.get("/api/sharepoint/status")
def sharepoint_status(db: Session = Depends(get_db)):
    row = db.query(AppSetting).filter(AppSetting.key == "sharepoint_config").first()
    if not row or not row.value:
        return {"configured": False, "next_steps": ["POST /api/sharepoint/config with tenant, app registration, site URL and library ID"]}
    cfg = json.loads(row.value)
    return {
        "configured": True,
        "site_url": cfg.get("site_url"),
        "library_id": cfg.get("library_id"),
        "acl_model": "Microsoft Graph + delegated or application permissions; enforce SharePoint ACLs on each file fetch (not implemented in this MVP).",
        "sync": "Use POST /api/sharepoint/sync to enqueue a sync job (stub).",
    }


@app.post("/api/sharepoint/sync")
def sync_sharepoint_documents(payload: SharePointSyncRequest, db: Session = Depends(get_db)):
    row = db.query(AppSetting).filter(AppSetting.key == "sharepoint_config").first()
    if not row or not row.value:
        raise HTTPException(status_code=400, detail="Configure SharePoint first via POST /api/sharepoint/config")
    return {
        "job_id": "sp-sync-stub-001",
        "status": "accepted_stub",
        "message": "Full SharePoint library sync with version mapping and ACL enforcement requires Microsoft Graph integration; this endpoint records intent only.",
        "requested": {"site_url": payload.site_url, "library_name": payload.library_name},
    }

@app.post("/api/auth/register")
def register_user(user: UserCreate, db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.username == user.username).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Username already registered")
    
    hashed_password = get_password_hash(user.password)
    new_user = User(username=user.username, password_hash=hashed_password, role=user.role)
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    return {"message": "User created successfully", "username": new_user.username, "role": new_user.role}

@app.post("/api/auth/login")
def login_user(user: UserLogin, db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.username == user.username).first()
    if not db_user or not verify_password(user.password, db_user.password_hash):
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid username or password")
    token = create_access_token({"sub": db_user.username, "role": db_user.role, "site": db_user.site or "default"})
    return {"message": "Login successful", "username": db_user.username, "role": db_user.role, "token": token}

# --- User Management Endpoints ---

@app.get("/api/users")
def get_users(db: Session = Depends(get_db), _user: dict = Depends(require_admin)):  # Fix #2
    users = db.query(User).all()
    return [{"id": u.id, "username": u.username, "role": u.role} for u in users]

@app.post("/api/users")
def create_user(user: UserCreate, db: Session = Depends(get_db), _admin: dict = Depends(require_admin)):  # Fix #2
    if db.query(User).filter(User.username == user.username).first():
        raise HTTPException(status_code=400, detail="Username already exists")
    hashed_password = get_password_hash(user.password)
    new_user = User(username=user.username, password_hash=hashed_password, role=user.role)
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    return {"id": new_user.id, "username": new_user.username, "role": new_user.role}

@app.delete("/api/users/{user_id}")
def delete_user(user_id: int, db: Session = Depends(get_db), _admin: dict = Depends(require_admin)):  # Fix #2
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    if user.username == "admin":
        raise HTTPException(status_code=400, detail="Cannot delete default admin")
    db.delete(user)
    db.commit()
    return {"message": "User deleted"}

# --- Document Management Endpoints ---

@app.get("/api/documents")
def get_documents(db: Session = Depends(get_db), _user: dict = Depends(require_auth)):  # Fix #3
    docs = db.query(DocumentMetadata).all()
    return docs

@app.post("/api/documents")
def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    doc_type: str = Form("Procédure"),
    criticality: str = Form("Medium"),
    version: str = Form("1.0"),
    owner: str = Form("QMS"),
    language: str = Form("fr"),
    site: str = Form("default"),
    db: Session = Depends(get_db)
):
    _, ext = os.path.splitext(file.filename)
    if ext.lower() not in SUPPORTED_EXTENSIONS:
        allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: {allowed}"
        )

    file_path = f"uploads/{file.filename}"
    with open(file_path, "wb") as buffer:
        buffer.write(file.file.read())
        
    new_doc = DocumentMetadata(
        filename=file.filename,
        file_path=file_path,
        doc_type=doc_type,
        criticality=criticality,
        version=version,
        owner=owner,
        language=language,
        site=site or "default",
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)
    
    # Ingest document into vector store in background
    metadata = {
        "filename": new_doc.filename,
        "doc_type": new_doc.doc_type,
        "criticality": new_doc.criticality,
        "version": new_doc.version,
        "owner": new_doc.owner,
        "language": new_doc.language,
        "site": new_doc.site,
        "uploaded_at": new_doc.uploaded_at.isoformat() if new_doc.uploaded_at else "",
    }
    background_tasks.add_task(ingest_document, file_path, new_doc.id, metadata)
    _log_activity(db, action="upload", username="admin", query=file.filename, response_summary=f"doc_id={new_doc.id} type={doc_type}")
    return new_doc

@app.delete("/api/documents/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(DocumentMetadata).filter(DocumentMetadata.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Try to remove file
    if os.path.exists(doc.file_path):
        try:
            os.remove(doc.file_path)
        except Exception as e:
            pass # ignore if file is already deleted or in use
            
    # Remove from vector index
    try:
        remove_document_from_index(doc_id)
    except Exception as e:
        print(f"Failed to remove doc from vector index: {e}")
        
    db.delete(doc)
    db.commit()
    _log_activity(db, action="delete", username="admin", query=doc.filename, response_summary=f"doc_id={doc_id} deleted")
    return {"message": "Document deleted"}


# --- Activity Logs Endpoint ---

@app.get("/api/logs")
def get_activity_logs(
    limit: int = 50,  # Fix #11: défaut réduit, pagination ajoutée
    page: int = 1,  # Fix #11: pagination
    action: str | None = None,
    username: str | None = None,
    db: Session = Depends(get_db),
    _user: dict = Depends(require_admin),  # Fix #2
):
    # Fix #11: pagination
    page = max(1, page)
    per_page = max(1, min(limit, 100))
    offset = (page - 1) * per_page
    q = db.query(ActivityLog)
    if action:
        q = q.filter(ActivityLog.action == action)
    if username:
        q = q.filter(ActivityLog.username == username)
    total = q.count()
    logs = q.order_by(ActivityLog.created_at.desc()).offset(offset).limit(per_page).all()
    return {
        "total": total,
        "page": page,
        "per_page": per_page,
        "pages": max(1, (total + per_page - 1) // per_page),
        "items": [
            {
                "id": log.id,
                "username": log.username,
                "action": log.action,
                "query": log.query,
                "document_ids": log.document_ids,
                "confidence": log.confidence,
                "confidence_score": log.confidence_score,
                "language_mode": log.language_mode,
                "response_summary": log.response_summary,
                "timestamp": log.created_at.isoformat() if log.created_at else None,
            }
            for log in logs
        ]
    }


# --- Chat Sessions (persistent, DB-backed) ---

class SessionSaveRequest(BaseModel):
    session_id: str
    title: str
    messages: list
    username: str

@app.get("/api/sessions")
def get_sessions(username: str, db: Session = Depends(get_db)):
    rows = db.query(ChatSession).filter(ChatSession.username == username).order_by(ChatSession.updated_at.desc()).limit(30).all()
    return [{"id": r.id, "title": r.title, "messages": json.loads(r.messages_json or "[]"), "updatedAt": r.updated_at.isoformat()} for r in rows]

@app.post("/api/sessions")
def save_session(payload: SessionSaveRequest, db: Session = Depends(get_db)):
    row = db.query(ChatSession).filter(ChatSession.id == payload.session_id).first()
    if row:
        row.title = payload.title[:80]
        row.messages_json = json.dumps(payload.messages)
        row.updated_at = datetime.now(timezone.utc).replace(tzinfo=None)
    else:
        row = ChatSession(
            id=payload.session_id,
            username=payload.username,
            title=payload.title[:80],
            messages_json=json.dumps(payload.messages),
            updated_at=datetime.now(timezone.utc).replace(tzinfo=None),
        )
        db.add(row)
    db.commit()
    return {"ok": True}

@app.delete("/api/sessions/{session_id}")
def delete_session(session_id: str, db: Session = Depends(get_db)):
    row = db.query(ChatSession).filter(ChatSession.id == session_id).first()
    if row:
        db.delete(row)
        db.commit()
    return {"ok": True}


# --- PFMEA Excel Export ---

@app.post("/api/generate/pfmea/export")
def export_pfmea_excel(payload: dict, db: Session = Depends(get_db)):
    """Export PFMEA rows as a formatted Excel file."""
    try:
        import openpyxl
        from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    except ImportError:
        raise HTTPException(status_code=500, detail="openpyxl not installed")

    rows = payload.get("rows", [])
    process = payload.get("process", "process")
    product = payload.get("product", "product")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "PFMEA"

    headers = ["#", "Étape process", "Produit", "Mode de défaillance", "Effets", "S", "O", "D", "RPN", "Actions recommandées"]
    header_fill = PatternFill(start_color="1E293B", end_color="1E293B", fill_type="solid")
    header_font = Font(color="F1F5F9", bold=True, size=10)
    thin = Side(style="thin", color="334155")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border

    ws.row_dimensions[1].height = 24
    col_widths = [5, 22, 16, 28, 28, 5, 5, 5, 7, 36]
    for i, w in enumerate(col_widths, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    red_fill   = PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid")
    amber_fill = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid")
    green_fill = PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid")

    for ri, row in enumerate(rows, start=2):
        rpn = int(row.get("rpn", 0) or 0)
        row_fill = red_fill if rpn > 200 else amber_fill if rpn > 100 else green_fill
        vals = [row.get("line", ri-1), row.get("process_step", ""), row.get("product", ""),
                row.get("failure_mode", ""), row.get("effects", ""),
                row.get("severity", ""), row.get("occurrence", ""), row.get("detection", ""),
                row.get("rpn", ""), row.get("recommended_actions", "")]
        for col, val in enumerate(vals, start=1):
            cell = ws.cell(row=ri, column=col, value=val)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = border
            if col == 9:  # RPN column
                cell.fill = row_fill
                cell.font = Font(bold=True)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    filename = f"pfmea_{process}_{product}.xlsx".replace(" ", "_")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# Fix #12 — Export PFMEA en PDF
@app.post("/api/generate/pfmea/export/pdf")
def export_pfmea_pdf(payload: dict):
    """Export PFMEA rows as a formatted PDF file."""
    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(status_code=500, detail="fpdf2 non installé")

    rows = payload.get("rows", [])
    process_name = payload.get("process", "Process")
    product_name = payload.get("product", "Product")

    pdf = FPDF(orientation="L", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()

    # Titre
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_fill_color(30, 41, 59)
    pdf.set_text_color(241, 245, 249)
    pdf.cell(0, 10, f"PFMEA — {process_name} / {product_name}", ln=True, fill=True, align="C")
    pdf.ln(4)

    # En-têtes colonnes
    headers = ["#", "Etape process", "Mode defaillance", "Effets", "S", "O", "D", "RPN", "Actions"]
    col_w   = [10, 38, 48, 48, 8, 8, 8, 12, 48]
    pdf.set_font("Helvetica", "B", 8)
    pdf.set_fill_color(51, 65, 85)
    pdf.set_text_color(203, 213, 225)
    for h, w in zip(headers, col_w):
        pdf.cell(w, 8, h, border=1, fill=True, align="C")
    pdf.ln()

    # Lignes
    pdf.set_font("Helvetica", "", 7.5)
    for row in rows:
        rpn_val = int(row.get("rpn", 0) or 0)
        if rpn_val > 200:
            pdf.set_fill_color(254, 226, 226)   # rouge
        elif rpn_val > 100:
            pdf.set_fill_color(254, 243, 199)   # ambre
        else:
            pdf.set_fill_color(220, 252, 231)   # vert
        pdf.set_text_color(15, 23, 42)
        vals = [
            str(row.get("line", "")), row.get("process_step", ""),
            row.get("failure_mode", ""), row.get("effects", ""),
            str(row.get("severity", "")), str(row.get("occurrence", "")),
            str(row.get("detection", "")), str(row.get("rpn", "")),
            row.get("recommended_actions", ""),
        ]
        row_h = 7
        for val, w in zip(vals, col_w):
            pdf.cell(w, row_h, str(val)[:60], border=1, fill=True)
        pdf.ln()

    raw = pdf.output()
    if isinstance(raw, str):
        raw = raw.encode("latin-1")
    buf = io.BytesIO(raw)
    buf.seek(0)
    fname = f"pfmea_{process_name}_{product_name}.pdf".replace(" ", "_")
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# Fix #13 — Checklist audit interactive (sauvegarde en DB)
class AuditChecklistSaveRequest(BaseModel):
    standard: str
    process: str
    checklist: list  # [{question: str, checked: bool, note: str}]
    username: str

@app.post("/api/audit/checklist/save")
def save_audit_checklist(
    payload: AuditChecklistSaveRequest,
    db: Session = Depends(get_db),
    _user: dict = Depends(require_auth),
):
    """Sauvegarde une checklist audit remplie par un utilisateur."""
    result = AuditResult(
        username=payload.username,
        standard=payload.standard,
        process=payload.process,
        checklist_json=json.dumps(payload.checklist, ensure_ascii=False),
    )
    db.add(result)
    db.commit()
    db.refresh(result)
    return {"id": result.id, "message": "Checklist sauvegardée"}

@app.get("/api/audit/checklist/history")
def get_audit_history(
    username: str | None = None,
    db: Session = Depends(get_db),
    _user: dict = Depends(require_auth),
):
    """Récupère l'historique des checklists audit."""
    q = db.query(AuditResult)
    if username:
        q = q.filter(AuditResult.username == username)
    results = q.order_by(AuditResult.created_at.desc()).limit(20).all()
    return [
        {
            "id": r.id,
            "username": r.username,
            "standard": r.standard,
            "process": r.process,
            "checklist": json.loads(r.checklist_json or "[]"),
            "created_at": r.created_at.isoformat() if r.created_at else None,
        }
        for r in results
    ]


# --- Stats Dashboard ---

@app.get("/api/stats")
def get_stats(db: Session = Depends(get_db)):
    """KPI stats for admin dashboard."""
    from sqlalchemy import func
    total_docs = db.query(DocumentMetadata).count()
    total_users = db.query(User).count()
    total_queries = db.query(ActivityLog).filter(ActivityLog.action == "chat").count()

    # Queries per day (last 7 days)
    seven_days_ago = datetime.now(timezone.utc).replace(tzinfo=None)
    from datetime import timedelta
    seven_days_ago = seven_days_ago - timedelta(days=7)
    recent_logs = db.query(ActivityLog).filter(
        ActivityLog.action == "chat",
        ActivityLog.created_at >= seven_days_ago
    ).all()

    daily: dict = {}
    for log in recent_logs:
        day = log.created_at.strftime("%Y-%m-%d") if log.created_at else "unknown"
        daily[day] = daily.get(day, 0) + 1

    # Confidence distribution
    conf_dist: dict = {}
    conf_logs = db.query(ActivityLog.confidence, func.count()).filter(
        ActivityLog.action == "chat", ActivityLog.confidence.isnot(None)
    ).group_by(ActivityLog.confidence).all()
    for label, cnt in conf_logs:
        conf_dist[label] = cnt

    # Top documents used
    doc_usage: dict = {}
    logs_with_docs = db.query(ActivityLog.document_ids).filter(ActivityLog.document_ids.isnot(None)).limit(200).all()
    for (doc_ids,) in logs_with_docs:
        for did in (doc_ids or "").split(","):
            did = did.strip()
            if did:
                doc_usage[did] = doc_usage.get(did, 0) + 1
    top_docs_raw = sorted(doc_usage.items(), key=lambda x: x[1], reverse=True)[:5]
    top_docs = []
    for did, cnt in top_docs_raw:
        doc = db.query(DocumentMetadata).filter(DocumentMetadata.id == int(did)).first() if did.isdigit() else None
        top_docs.append({"doc_id": did, "filename": doc.filename if doc else did, "count": cnt})

    return {
        "total_documents": total_docs,
        "total_users": total_users,
        "total_queries": total_queries,
        "queries_per_day": daily,
        "confidence_distribution": conf_dist,
        "top_documents": top_docs,
    }


# --- Ollama Models List ---

@app.get("/api/ollama/models")
def get_ollama_models(db: Session = Depends(get_db)):
    """Fetch available models from local Ollama instance."""
    import urllib.request, urllib.error
    cfg = db.query(LLMConfig).filter(LLMConfig.provider == "ollama").first()
    host = (cfg.base_url.strip() if cfg and cfg.base_url else "") or os.getenv("OLLAMA_BASE_URL", "http://127.0.0.1:11434")
    try:
        with urllib.request.urlopen(f"{host}/api/tags", timeout=3) as resp:
            data = json.loads(resp.read())
            return {"models": [m["name"] for m in data.get("models", [])]}
    except Exception:
        return {"models": []}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
